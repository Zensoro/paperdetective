"""Phase A input ingestion: PDF/Word/image/text -> Document."""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path

from PIL import Image


@dataclass
class Document:
    paper_id: str
    text: str = ""
    images: list = field(default_factory=list)  # [(image_id, PIL.Image)]
    tables: list = field(default_factory=list)  # [dict]

    def image_ids(self) -> list[str]:
        return [i[0] for i in self.images]


def ingest_text(text: str, paper_id: str = "doc") -> Document:
    return Document(paper_id=paper_id, text=text)


def extract_images(data, format: str = "png", prefix: str = "img") -> list[tuple[str, Image.Image]]:
    with Image.open(data) as img:
        return [(f"{prefix}0", img.copy())]


# ---------------------------------------------------------------------------
# PDF image extraction + page-furniture filtering
#
# Journal "printable" PDFs (e.g. PLoS) often embed, on EVERY page, the SAME
# full-page raster XObject (a 2550x3300 placeholder). Without filtering, those
# page rasters are treated as figures and pHash reports "every image is a
# duplicate of every other image" (a 100+ finding false-positive cascade, first
# observed on a real ORI-confirmed fabrication case — see
# docs/case-studies/her3-brand-2013.md).
#
# Two O(1)-per-image heuristics remove >=90% of such false positives:
#   1. document-level hash dedup: an image whose fingerprint appears on >=80%
#      of pages is page furniture (watermark / header / rasterized page);
#   2. size-band filter: >4000px (full-page raster) or <100px (logo / form
#      field) images are not substantive figures.
# ---------------------------------------------------------------------------
PAGE_FURNITURE_RATIO = 0.8
MAX_FIGURE_DIM = 4000
MIN_FIGURE_DIM = 100


def _thumb_fingerprint(img: Image.Image, size: int = 16) -> bytes:
    """Content fingerprint: 16x16 grayscale downsample (fast, order-invariant)."""
    return img.convert("L").resize((size, size)).tobytes()


def _page_key(img_id: str) -> str:
    """Map an image id (e.g. 'page3_X0.png') to its page key (e.g. 'page3')."""
    import re
    m = re.match(r"(page\d+)", img_id)
    return m.group(1) if m else img_id


def _filter_furniture(
    images: list[tuple[str, Image.Image]],
    furniture_ratio: float = PAGE_FURNITURE_RATIO,
    max_dim: int = MAX_FIGURE_DIM,
    min_dim: int = MIN_FIGURE_DIM,
) -> list[tuple[str, Image.Image]]:
    """Drop shared page-furniture / oversized / tiny images before detection.

    Furniture is judged at PAGE granularity: an image whose content fingerprint
    appears on >= `furniture_ratio` of the pages is page furniture (watermark,
    header, rasterized page placeholder), regardless of how many total images
    the document yields.
    """
    if not images:
        return images
    pages = sorted({_page_key(i) for i, _ in images})
    n_pages = max(1, len(pages))
    from collections import Counter
    page_counts: Counter[str] = Counter()
    for img_id, img in images:
        fp = _thumb_fingerprint(img)
        page_counts[(fp, _page_key(img_id))] += 1  # per page, per fingerprint
    # fingerprints appearing on >= ratio of distinct pages are furniture
    fp_pages: Counter[str] = Counter()
    for (fp, page) in page_counts:
        fp_pages[fp] += 1
    kept = []
    for img_id, img in images:
        if fp_pages[_thumb_fingerprint(img)] >= n_pages * furniture_ratio:
            continue  # same content on almost every page -> page furniture
        w, h = img.size
        if max(w, h) > max_dim or min(w, h) < min_dim:
            continue
        kept.append((img_id, img))
    return kept


def _extract_pdf_images(reader, prefix: str = "page") -> list[tuple[str, Image.Image]]:
    """Extract embedded images from a pypdf PdfReader (best-effort per page)."""
    out: list[tuple[str, Image.Image]] = []
    for pi, page in enumerate(reader.pages):
        try:
            for im in page.images:
                name = getattr(im, "name", None) or f"X{len(out)}"
                try:
                    out.append((f"{prefix}{pi + 1}_{name}", Image.open(io.BytesIO(im.data)).convert("RGB")))
                except Exception:
                    continue  # malformed image object: skip silently
        except Exception:
            continue  # page without extractable images
    return out


def ingest_path(path: str) -> Document:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".txt":
        return ingest_text(p.read_text(encoding="utf-8", errors="ignore"), paper_id=p.stem)
    if suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
        with Image.open(p) as img:
            return Document(paper_id=p.stem, images=[(p.stem, img.copy())])
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            return Document(paper_id=p.stem, text="[PDF support requires: pip install paperdetective[pdf]]")
        reader = PdfReader(str(p))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        images = _filter_furniture(_extract_pdf_images(reader))
        return Document(paper_id=p.stem, text=text, images=images)
    if suffix == ".docx":
        try:
            import docx
        except ImportError:
            return Document(paper_id=p.stem, text="[Word support requires: pip install paperdetective[docx]]")
        d = docx.Document(str(p))
        text = "\n".join(par.text for par in d.paragraphs)
        return Document(paper_id=p.stem, text=text)
    if suffix == ".doc":
        raise ValueError("legacy .doc is not supported; please convert to .docx first")
    raise ValueError(f"unsupported file type: {suffix}")
